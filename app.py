from flask import Flask, render_template, request, send_file
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import time
from werkzeug.utils import secure_filename

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        pdf = request.files.get("pdf")

        # Validate PDF file
        if not pdf or not pdf.filename.lower().endswith(".pdf"):
            return render_template("index.html", error="Only PDF files are allowed.")

        # Validate file size
        pdf.seek(0, os.SEEK_END)
        if pdf.tell() > MAX_FILE_SIZE:
            return render_template("index.html", error="File size must be less than 5MB.")
        pdf.seek(0)

        filename = secure_filename(pdf.filename)
        pdf_path = os.path.join(UPLOAD_FOLDER, filename)
        pdf.save(pdf_path)

        # Extract text from PDF
        pdf_doc = fitz.open(pdf_path)
        extracted_text = ""
        for page in pdf_doc:
            extracted_text += page.get_text()

        # Create Word document
        output_filename = f"converted_{int(time.time())}.docx"
        word_path = os.path.join(OUTPUT_FOLDER, output_filename)

        doc = Document()

        for line in extracted_text.split("\n"):
            line = line.strip()

            # Empty line
            if not line:
                doc.add_paragraph("")
                continue

            # Detect headings (ALL CAPS)
            if line.isupper():
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(14)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Normal paragraph
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)

        doc.save(word_path)

        return render_template(
            "index.html",
            success=True,
            download_file=output_filename
        )

    return render_template("index.html")


@app.route("/download/<filename>")
def download(filename):
    return send_file(
        os.path.join(OUTPUT_FOLDER, filename),
        as_attachment=True
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)

